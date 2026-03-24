#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MD转DOCX工具 - AI投资Agent
支持Markdown表格转换为Word表格，使用楷体字体
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.enum.table import Wd_Table_Alignment
except ImportError:
    print("请先安装python-docx: pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name='楷体', size=11):
    """设置文字字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def parse_md_table(lines, start_idx):
    """解析Markdown表格"""
    rows = []
    i = start_idx

    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # 跳过分隔行
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        # 解析数据行
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1

    return rows, i


def add_table_to_doc(doc, rows, font_name='楷体'):
    """添加表格到Word文档"""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = Wd_Table_Alignment.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = ''
                run = cell.paragraphs[0].add_run(cell_text)
                set_run_font(run, font_name, 10 if i == 0 else 10)
                if i == 0:  # 表头加粗
                    run.bold = True


def md_to_docx(md_content, output_path, font_name='楷体'):
    """将Markdown内容转换为Word文档"""
    doc = Document()
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 跳过分隔线
        if stripped == '---':
            i += 1
            continue

        # 处理标题
        if stripped.startswith('#'):
            level = len(stripped.split()[0])
            text = stripped.lstrip('#').strip()

            if level == 1:
                para = doc.add_heading(text, level=0)
            elif level <= 4:
                para = doc.add_heading(text, level=level)
            else:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                set_run_font(run, font_name, 11)

            i += 1
            continue

        # 处理表格
        if stripped.startswith('|'):
            rows, i = parse_md_table(lines, i)
            if rows:
                add_table_to_doc(doc, rows, font_name)
                doc.add_paragraph()
            continue

        # 处理列表
        if stripped.startswith('- ') or stripped.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            text = stripped[2:]
            run = para.add_run(text)
            set_run_font(run, font_name)
            i += 1
            continue

        # 处理数字列表
        if re.match(r'^\d+\.\s', stripped):
            para = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', stripped)
            run = para.add_run(text)
            set_run_font(run, font_name)
            i += 1
            continue

        # 处理普通段落
        para = doc.add_paragraph()

        # 处理加粗和斜体
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                run = para.add_run(part)
            set_run_font(run, font_name)

        i += 1

    doc.save(output_path)
    print(f"已保存: {output_path}")


def convert_file(md_path, docx_path=None, font_name='楷体'):
    """转换单个文件"""
    md_path = Path(md_path)
    if not md_path.exists():
        print(f"文件不存在: {md_path}")
        return False

    if docx_path is None:
        docx_path = md_path.with_suffix('.docx')

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    md_to_docx(content, str(docx_path), font_name)
    return True


def main():
    """命令行入口"""
    if len(sys.argv) < 2:
        print("用法: python md2docx.py <markdown文件> [输出文件] [字体名称]")
        print("示例: python md2docx.py report.md report.docx 楷体")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) > 2 else None
    font_name = sys.argv[3] if len(sys.argv) > 3 else '楷体'

    convert_file(md_path, docx_path, font_name)


if __name__ == '__main__':
    main()
